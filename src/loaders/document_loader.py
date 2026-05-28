import os
from langchain_core.documents import Document

import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# =====================================================
# TXT LOADER
# =====================================================
class TXTLoader:
    def load(self, path):
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except:
            return ""


# =====================================================
# DOCX LOADER
# =====================================================
class DOCXLoader:
    def load(self, path):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            return "\n".join([p.text for p in doc.paragraphs])
        except:
            return ""


# =====================================================
# CSV LOADER (IMPROVED)
# =====================================================
class CSVLoader:
    def load(self, path):
        try:
            import pandas as pd
            df = pd.read_csv(path)
            return df.to_string(index=False)
        except:
            return ""


# =====================================================
# PDF LOADER (FIXED: TEXT + TABLE + OCR)
# =====================================================
class PDFLoader:

    def load(self, path):

        text = ""

        try:
            import fitz  # PyMuPDF
            import pdfplumber
            from pdf2image import convert_from_path

            # =========================================
            # 1. TEXT EXTRACTION (PyMuPDF)
            # =========================================
            doc = fitz.open(path)

            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text:
                    text += f"\nPAGE_{page_num + 1}:\n{page_text}"

            # =========================================
            # 2. TABLE EXTRACTION (STRUCTURED)
            # =========================================
            try:
                with pdfplumber.open(path) as pdf:
                    for page_num, page in enumerate(pdf.pages):

                        tables = page.extract_tables()

                        for table in tables:
                            if table:
                                text += f"\n\nTABLE_PAGE_{page_num + 1}:\n"

                                for row in table:
                                    clean_row = [
                                        str(cell).strip() if cell else ""
                                        for cell in row
                                    ]
                                    text += " | ".join(clean_row) + "\n"

            except Exception as e:
                print("TABLE ERROR:", e)

            # =========================================
            # 3. OCR FALLBACK (ONLY IF LOW TEXT)
            # =========================================
            if len(text.strip()) < 200:

                try:
                    images = convert_from_path(path)

                    ocr_text = ""

                    for img in images:
                        ocr_text += pytesseract.image_to_string(img)

                    text += "\nOCR_TEXT:\n" + ocr_text

                except Exception as e:
                    print("OCR ERROR:", e)

        except Exception as e:
            print("PDF LOAD ERROR:", e)
            return ""

        return text


# =====================================================
# GET LOADER
# =====================================================
def get_loader(ext):
    return {
        ".pdf": PDFLoader(),
        ".txt": TXTLoader(),
        ".docx": DOCXLoader(),
        ".csv": CSVLoader()
    }.get(ext)


# =====================================================
# PROCESS FILES (FIXED + SAFE + CLEAN DOCS)
# =====================================================
def process_uploaded_files(files):

    documents = []
    os.makedirs("data/uploads", exist_ok=True)

    for file in files:

        try:
            file_bytes = file.getvalue()
            path = os.path.join("data/uploads", file.name)

            with open(path, "wb") as f:
                f.write(file_bytes)

            ext = os.path.splitext(path)[1].lower()

            loader = get_loader(ext)

            if not loader:
                continue

            content = loader.load(path)

            if not content:
                continue

            content = str(content).strip()

            if len(content) < 10:
                continue

            # =========================================
            # IMPORTANT: BETTER METADATA FOR RETRIEVAL
            # =========================================
            documents.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": file.name,
                        "type": ext.replace(".", ""),
                        "length": len(content)
                    }
                )
            )

        except Exception as e:
            print("FILE ERROR:", file.name, e)

    print("[DEBUG] FINAL DOC COUNT:", len(documents))

    return documents